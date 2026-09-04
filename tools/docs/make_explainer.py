"""Generate the SignStream project explainer / defence guide as a .docx.

    pip install python-docx
    python tools/docs/make_explainer.py docs/SignStream-Explained.docx

The document is generated rather than hand-written so its numbers can be
re-checked against the repository and regenerated when they change. Every
figure in it was measured, not estimated.
"""

from __future__ import annotations

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x55, 0x5F, 0x67)
CODE_BG = "F2F4F6"
NOTE_BG = "FFF6E5"

doc = Document()

# ── Page + base styles ───────────────────────────────────────────────────────
for s in doc.sections:
    s.top_margin = Inches(0.9)
    s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

for name, size, color in (("Heading 1", 18, ACCENT), ("Heading 2", 14, ACCENT), ("Heading 3", 12, ACCENT)):
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(14 if name != "Heading 1" else 20)
    st.paragraph_format.space_after = Pt(6)


def shade(par, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    par._p.get_or_add_pPr().append(el)


def code(text, fill=CODE_BG):
    for line in text.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.18)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        shade(p, fill)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def para(text, bold_prefix=None, style=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix:
        p.add_run(bold_prefix).bold = True
    p.add_run(text)
    return p


def bullet(text, bold_prefix=None):
    return para(text, bold_prefix, style="List Bullet")


def note(title, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(title + "  ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    shade(p, NOTE_BG)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def qa(question, answer):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Q.  " + question)
    r.bold = True
    r.font.color.rgb = ACCENT
    a = doc.add_paragraph()
    a.paragraph_format.left_indent = Inches(0.25)
    a.paragraph_format.space_after = Pt(10)
    a.add_run(answer)


# ── Title page ───────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(150)
r = t.add_run("SignStream")
r.font.size = Pt(40)
r.bold = True
r.font.color.rgb = ACCENT

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Real-time translation of streaming audio into sign language,\nrendered by a 3D avatar in the browser")
r.font.size = Pt(13)
r.font.color.rgb = MUTED

s2 = doc.add_paragraph()
s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
s2.paragraph_format.space_before = Pt(40)
r = s2.add_run("Project Explainer and Defence Guide")
r.font.size = Pt(12)
r.bold = True

s3 = doc.add_paragraph()
s3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s3.add_run("Kwame Nkrumah University of Science and Technology\nFinal-Year Project")
r.font.size = Pt(11)
r.font.color.rgb = MUTED

s4 = doc.add_paragraph()
s4.alignment = WD_ALIGN_PARAGRAPH.CENTER
s4.paragraph_format.space_before = Pt(30)
r = s4.add_run("Languages: American Sign Language and Ghanaian Sign Language\n"
               "3,199 keypoint clips  ·  12 avatars  ·  MV3 extension + AWS serverless backend")
r.font.size = Pt(10)
r.font.color.rgb = MUTED

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ── How to use ───────────────────────────────────────────────────────────────
doc.add_heading("How to use this document", 1)
para("This is written to be read once, straight through, before a defence. It explains "
     "not only what SignStream does but why each decision was made, because a viva "
     "rewards the second far more than the first.")
para("Section 10 is a list of the questions an examiner is most likely to ask, with "
     "answers. If you are short of time, read sections 2, 5 and 10.")
note("On honesty:", "Every number in this document was measured from the repository, not "
     "estimated. Where something does not work, it says so. A limitation you raise "
     "yourself is a design boundary; the same limitation raised by an examiner is a gap "
     "you missed.")

# ── 1. The problem ───────────────────────────────────────────────────────────
doc.add_heading("1.  The problem", 1)
para("Ghana's 2021 Population and Housing Census records 470,737 people with some degree "
     "of hearing loss, of whom over 211,000 are estimated to be Deaf or hard of hearing. "
     "Ghanaian Sign Language is used by the majority of Deaf Ghanaians. Streaming video "
     "— lectures, news, entertainment — is effectively closed to them unless it "
     "is captioned, and even then it is only partly open.")
note("Be careful with this number.", "Sources disagree and the data is poorly documented. "
     "The Ghana National Association of the Deaf puts the figure near 110,000 (0.4% of the "
     "population); the Ghana Statistical Service reports 211,712. Cite the census, "
     "acknowledge the disagreement, and do not round upward — an inflated figure is the "
     "easiest thing in a viva to challenge.")

doc.add_heading("Captions are not the same as sign language", 2)
para("This is the point the whole project rests on, and it is the one most people get "
     "wrong. Sign languages are not visual encodings of spoken languages. GhSL has its "
     "own grammar, its own word order, and no written form. For someone whose first "
     "language is GhSL, English captions are a ")
p = doc.paragraphs[-1]
p.add_run("second language").bold = True
p.add_run(", read at a disadvantage, often at a reading level below hearing peers because "
          "literacy in a language you cannot hear is genuinely harder to acquire.")
para("A signed rendering is not a convenience on top of captions. For a signer it is the "
     "message itself.")

doc.add_heading("Why this is not already solved", 2)
bullet("Human interpreters do not scale to the volume of video published each day, and "
       "cost per hour is prohibitive for a Ghanaian audience.", "Interpreters cost. ")
bullet("Almost all sign-language AI research targets ASL and German Sign Language. GhSL "
       "is a low-resource language: this project established that the single public "
       "word-level GhSL keypoint dataset in existence contains 1,200 signs from one "
       "signer.", "The data does not exist. ")
bullet("Generating fluent signing from text is an unsolved research problem. Anything "
       "claiming otherwise at final-year scale is overclaiming.", "Generation is unsolved. ")

# ── 2. What was built ────────────────────────────────────────────────────────
doc.add_heading("2.  What was built", 1)
para("SignStream is a Chrome Manifest V3 browser extension. A Deaf user opens a video, "
     "and an avatar appears over the page performing the audio in sign language.")

table(["Property", "Decision"],
      [["Direction", "One-way: audio to sign. There is no sign-to-speech path."],
       ["Surface", "A browser extension overlaying a page the user already opened."],
       ["Languages", "ASL and GhSL. BSL is deliberately blocked."],
       ["Client compute", "None. All inference is cloud-side; the browser only renders."],
       ["Avatar", "Pre-recorded human keypoints replayed on a rigged 3D model."]],
      widths=[1.4, 5.0])

doc.add_heading("The scope was deliberately narrowed", 2)
para("The project began as something larger: a bidirectional GhSL-to-speech translator "
     "inside a video-conferencing desktop application. It was narrowed to the one-way "
     "browser extension that exists now.")
note("Defend this as judgement, not retreat.", "The bidirectional path requires sign "
     "recognition from live video, which is a harder research problem than the whole of "
     "the rest of the system. Narrowing produced something that works end to end instead "
     "of four things that each half-work. If you find old documents mentioning video "
     "conferencing or sign-to-text, they are from the superseded scope.")

doc.add_heading("Why BSL is blocked rather than offered", 2)
para("No public BSL keypoint dataset exists. Offering BSL in the language menu would show "
     "a Deaf user an avatar standing still — a promise the system cannot keep. It is "
     "blocked in the frontend list and in the backend allowlist so the two cannot drift "
     "apart. Refusing to offer a language is a feature.")

# ── 3. Architecture ──────────────────────────────────────────────────────────
doc.add_heading("3.  How it works, end to end", 1)
code("""
tab audio  (chrome.tabCapture)
    |   16 kHz mono PCM, 250 ms frames
    v
offscreen document ----WebSocket----> API Gateway
                                          |
                                          v
                              ws-audio-ingest  ->  asr  (Moonshine ONNX)
                                          |
                                          v
                              text-to-gloss
                              "thank you" -> THANK-YOU -> asl-thank-you-v1
                                          |
    +-------------- signId <---------------+
    v
service worker ----relay----> content script
                                  |  fetch <cdn>/asl/asl-thank-you-v1.json
                                  v
                        Three.js avatar plays the keypoint clip
""")

doc.add_heading("Stage by stage", 2)
table(["Stage", "What happens"],
      [["1. Capture", "chrome.tabCapture takes the tab's audio in an offscreen document. "
                      "Capturing removes audio from the speakers, so it is explicitly routed "
                      "back through a gain node the user controls."],
       ["2. Encode", "An AudioWorklet resamples to 16 kHz mono PCM in 250 ms frames. "
                     "AudioWorklet, not ScriptProcessorNode, which is deprecated."],
       ["3. Transport", "A WebSocket to API Gateway. Binary frames are audio; text frames "
                        "are control messages."],
       ["4. Recognise", "Moonshine ONNX on Lambda CPU produces partial and final transcripts."],
       ["5. Map", "Rule-based tokenisation, then greedy longest-match against a word-to-gloss "
                  "dictionary. Deterministic, no inference."],
       ["6. Identify", "Gloss becomes a canonical sign id: THANK-YOU -> asl-thank-you-v1."],
       ["7. Fetch", "The content script fetches that clip from CloudFront and caches it."],
       ["8. Render", "Keypoints are retargeted to the avatar's skeleton and played."]],
      widths=[1.1, 5.3])

doc.add_heading("Two different things are called “dictionary”", 2)
para("This confuses everyone who reads the code, so be ready for it:")
code("""
backend/functions/text-to-gloss/dictionaries/   word  -> gloss label
dictionary/                                     gloss -> motion clip
""")

doc.add_heading("The full lookup chain", 2)
code("""
spoken "hello"
  -> dictionaries/asl.json        English word -> gloss
  -> gloss HELLO
  -> mapper.to_sign_id()          gloss -> sign id
  -> asl-hello-v1
  -> dictionary/asl/hello-v1.json sign id -> keypoint clip
  -> avatar.ts                    clip -> avatar motion
""")

# ── 4. The avatar ────────────────────────────────────────────────────────────
doc.add_heading("4.  The avatar: how motion is produced", 1)
para("This is the most technically interesting part of the system and the most likely "
     "target for detailed questions.")

doc.add_heading("Clips store positions; skeletons need rotations", 2)
para("A clip records where a human signer's joints were: 67 tracked points per frame — "
     "25 body, 21 per hand. A rigged 3D model cannot be driven by positions. It is "
     "animated by rotating each bone relative to its parent. Every frame, the system must "
     "answer: what rotation makes this bone point from its start keypoint toward its end "
     "keypoint?")

note("This is not inverse kinematics.", "IK is for when you know only where the hand "
     "ended up and must infer the arm. Here every joint position is known, so each bone "
     "is solved directly. There is no solver to converge and nothing to oscillate. "
     "Roughly 40 bones per frame is a few hundred quaternion operations — negligible "
     "beside rendering.")

code("""
targetLocal      = inverse(parentWorldRotation) * targetWorld
bone.quaternion  = rotationFrom(boneRestAxis, targetLocal)
""")
para("Parents are solved before children, which is what makes a chain hang together: "
     "rotating the upper arm carries the forearm and the whole hand with it.")

doc.add_heading("Why the retargeting runs at runtime rather than being pre-baked", 2)
para("Baking rotations into the clips would roughly double their size, force all 3,199 to "
     "be regenerated and re-uploaded, invalidate the CDN cache — and lock every clip to "
     "one skeleton. Same mathematics, worse place for it.")

doc.add_heading("Anatomical limits are enforced, because the data is 2D", 2)
para("The source keypoints are two-dimensional: z is zero on every point. A finger pointing "
     "toward the camera projects to almost nothing, and normalising a near-zero vector turns "
     "tracker noise into an essentially random direction.")
para("Measured across roughly 120,000 joint-frames of the real dictionary, 6.4% of solved "
     "finger poses exceeded 110 degrees — a finger folding backwards through the palm. "
     "Three defences were added:")
bullet("Segments shorter than 0.012 m in clip space yield no direction at all, so the bone "
       "holds its previous rotation rather than snapping to noise.")
bullet("Finger joints are clamped to 110 degrees of flexion and 8 degrees of hyperextension.")
bullet("Joints past the knuckle are treated as pure hinges, since they physically cannot "
       "abduct. The flexion axis is measured from the rest pose rather than declared, so "
       "nothing rig-specific is assumed.")

doc.add_heading("Smoothing", 2)
para("Two exponential filters, frame-rate independent so a 144 Hz monitor does not smooth "
     "differently from a 60 Hz one: 70 ms for the body, 150 ms for the fingers. Signs "
     "cross-fade over 140 ms; letters over 45 ms, because a full hand-over between letters "
     "would visually punctuate a word into pieces.")
note("Measured, not asserted.", "An offline harness replays real clips through the shipped "
     "code and measures jerk — the frame-to-frame change in angular velocity. Across all 12 "
     "avatars the smoothing removes 90 to 91 percent of it, with 0 unmapped bones and no "
     "numerical failures.")

doc.add_heading("Why VRM rather than glTF", 2)
para("An earlier iteration used MakeHuman models exported through Blender to glTF. Each "
     "needed a hand-written table mapping keypoints to that model's particular bone names, "
     "discovered by reading an export. One dropped separator once caused 46 of 47 bones to "
     "fail silently — the avatar simply did not move correctly, with no error anywhere.")
para("A VRM file declares its own humanoid. Bone identity comes from the file, so the same "
     "map drives any VRM. Adding an avatar became a data change rather than a code change.")

# ── 5. Design decisions ──────────────────────────────────────────────────────
doc.add_heading("5.  Design decisions and their justifications", 1)
para("These are the questions worth rehearsing. Each is a real trade-off with a reason.")

doc.add_heading("Moonshine instead of Whisper", 2)
table(["Consideration", "Reasoning"],
      [["Deployment target", "AWS Lambda CPU on the free tier. No GPU budget."],
       ["Whisper's problem", "It pads audio to a fixed 30-second window and re-transcribes it, "
                             "so short utterances cost the same as long ones."],
       ["Moonshine", "Processes only the audio it is given. Purpose-built for low-latency CPU "
                     "inference. MIT licensed."],
       ["Parakeet TDT", "Higher accuracy, retained on paper as a future upgrade. Needs an "
                        "always-on GPU (~$440/month), which breaks the cost constraint."]],
      widths=[1.5, 4.9])
note("A real bug worth mentioning.", "Moonshine initially took 3 to 6 seconds per inference. "
     "The cause was calling transcribe(audio, \"model-name\"), which rebuilds the ONNX model "
     "on every call. Passing the model object instead gave a 6 to 12 times speed-up. "
     "This is a good answer to “what went wrong and how did you find it?”")

doc.add_heading("Pre-recorded clips instead of generated signing", 2)
para("Generating fluent signing from text is unsolved research. Replaying recorded human "
     "keypoints guarantees that every sign shown is a sign a real signer actually made. The "
     "cost is a bounded vocabulary; the benefit is that the system is never confidently wrong.")
para("For a Deaf viewer, a plausible-looking invented sign is worse than no sign, because it "
     "cannot be distinguished from a real one.")

doc.add_heading("Captions bypass the ASR entirely", 2)
para("When a video publishes a caption track, its text is taken directly with the media "
     "timestamp attached, and signs are scheduled for the moment the words are spoken rather "
     "than played on arrival. This gives true synchronisation on pre-recorded video, costs no "
     "inference, and is more accurate than any ASR. Moonshine runs only for live streams and "
     "unsubtitled video, which is where it is genuinely needed.")

doc.add_heading("Signs expire rather than queueing", 2)
para("Signing is slower than speech. A sign runs about 2.4 seconds; speech delivers two to "
     "three words per second. A queue that simply grows means the avatar is soon describing "
     "something said minutes ago — which contradicts what is on screen and is worse than "
     "showing nothing.")
para("Signs therefore carry the media time they belong to and are dropped past 6 seconds of "
     "age. This bounds the lag permanently instead of letting it accumulate. The avatar shows "
     "what it can, in time, rather than everything, late.")

doc.add_heading("Fingerspelling refuses partial alphabets", 2)
para("A word is spelled only if every one of its letters has a clip. Currently both languages "
     "have 20 of 26 letters, so nothing is spelled at all.")
note("Defend the refusal.", "A guessed handshape is not an approximate letter — it is a "
     "different letter, rendered confidently. Spelling AIRPODS with an invented A does not "
     "communicate the word; it communicates a wrong one. Partial spelling is not spelling.")
para("Spelled letters play at 2.2 times normal speed. This is not cosmetic: the mean letter "
     "clip is 1,244 ms, so a five-letter word at normal pace takes 6.22 seconds and expires "
     "against the 6-second limit before it finishes, arriving with letters missing. At 2.2x it "
     "takes 2.83 seconds, and even the 8-letter maximum fits in 4.5 seconds.")

doc.add_heading("The GhSL alphabet is the ASL alphabet", 2)
para("Deaf education in Ghana began in 1957 under Andrew Foster, an African-American "
     "missionary who brought ASL with him. GhSL grew from that root. Its lexicon is its own, "
     "but its manual alphabet is the ASL one: one-handed, 22 distinct handshapes covering 26 "
     "letters, with h/u, k/p and g/l sharing a shape distinguished by movement or orientation.")
para("So the ASL letter clips were copied to GhSL. This is linguistics, not a shortcut — but "
     "each copied clip records derivedFrom and derivedNote in its own JSON so that nobody "
     "later mistakes them for recordings by a Ghanaian signer. Lexical signs are never shared "
     "this way.")

doc.add_heading("No AI on the client", 2)
para("The target user may be on a low-specification laptop. Running inference in the browser "
     "would exclude exactly the users the project exists for. The browser renders; the cloud "
     "thinks.")

# ── 6. Engineering ───────────────────────────────────────────────────────────
doc.add_heading("6.  Engineering detail worth knowing", 1)

doc.add_heading("Manifest V3 service workers are ephemeral", 2)
para("An MV3 service worker is killed after roughly 30 seconds idle. Any state held in a "
     "module variable is lost, and the worker restarts believing nothing is happening. This "
     "caused duplicate captures and a silent stream: two captures of the same tab, the second "
     "delivering nothing.")
para("Capture state therefore lives in chrome.storage.session, which survives worker restarts, "
     "stays in memory only, and is cleared when the browser closes — so a crash can never "
     "leave a stale “capturing” flag on disk.")

doc.add_heading("The overlay must be tab-isolated", 2)
para("The content script runs on every page. On load it asks the service worker whether "
     "capture is running, and mounts the avatar if so. The worker answered with the "
     "extension-wide flag — which is a property of the extension, not of a page — so opening "
     "any site while a video played in another tab put a signing avatar on it.")
para("The content script cannot filter this itself, because a page has no way to learn its own "
     "tab id. The answer is narrowed in the service worker, which knows it from sender.tab.id. "
     "A regression test pins all seven cases.")

doc.add_heading("Resampling drift", 2)
para("The AudioWorklet resampler produced 2.9% too much audio at 96 kHz because of an "
     "unbounded buffer index. It was found by bundling the shipped worklet and running it in "
     "Node against synthetic input — not by listening to it.")

doc.add_heading("A verification approach that repeatedly found real bugs", 2)
para("Several harnesses under tools/rt bundle the actual shipped TypeScript with esbuild and "
     "run it in Node against the real avatar files and the real clips. This caught the "
     "retargeting errors, the resampler drift, and the tab-isolation leak.")
note("A caution that is itself a good answer.", "The bundle is a compiled copy and does not "
     "update itself. It once sat three weeks stale against a rig set that had been replaced "
     "entirely — every rig in it deleted — and still reported success. A stale harness is "
     "worse than no harness, because it passes while testing code that is no longer shipped.")

doc.add_heading("Testing", 2)
table(["Suite", "Coverage"],
      [["backend/scripts/test-all.sh", "7 suites across all Lambdas and the shared layer."],
       ["pose-generator", "Keypoint conversion from OpenPose and WLASL."],
       ["tools/rt/smooth.mjs", "Bone resolution, numerical health and jerk, all 12 avatars."],
       ["tools/rt/tabs.mjs", "Overlay tab isolation, 7 cases."],
       ["dev/docker-compose.yml", "The whole pipeline on a laptop with the real modules."]],
      widths=[2.2, 4.2])
para("The local stack replaces only the AWS transport. The ASR engine, the normaliser and the "
     "gloss mapper are the production code, so what the avatar signs is never simulated.")

# ── 7. Data ──────────────────────────────────────────────────────────────────
doc.add_heading("7.  The data, and its constraints", 1)
table(["Dictionary", "Size", "Source", "Licence"],
      [["GhSL", "1,198 signs + 20 letters",
        "Ghanaian Sign Language Lexicon (Fragkiadakis, Nyst & Nyarko, 2021)", "CC BY 4.0"],
       ["ASL", "1,981 signs + 20 letters", "WLASL (Li et al., 2020)",
        "C-UDA — non-commercial"],
       ["Avatars", "12 VRM models", "Polygonal Mind, 100 Avatars", "CC0"]],
      widths=[0.9, 1.7, 2.5, 1.3])

note("Know this before you are asked.", "The ASL dictionary may not be used commercially. "
     "C-UDA permits academic and computational use only. GhSL is CC BY 4.0 and carries no such "
     "limit, so the Ghana-facing deliverable — the one the project is actually for — "
     "is clear. Commercialising ASL would require replacing its clips with a permissively "
     "licensed source.")

doc.add_heading("A finding worth presenting as a result", 2)
para("A deliberate search established that the Ghanaian Sign Language Lexicon is, as far as "
     "can be determined, the only public word-level GhSL keypoint dataset in existence: 1,200 "
     "signs, one signer. Larger GhSL resources exist but do not fit — SignTalk-GH is 10,000 "
     "sentence-level videos with no keypoints, and AfriSign is licensed no-derivatives.")
para("This is not a gap in the project. It is a finding about the field, and it explains why "
     "the vocabulary is the size it is. Low-resource languages are low-resource because this "
     "work is rare and unfunded, not because the languages are simple.")

doc.add_heading("Quality control on the ASL conversion", 2)
para("WLASL is scraped from public video of wildly varying quality, and on blurred or "
     "low-resolution clips the pose tracker returns all-zero hand keypoints — producing a clip "
     "where the arms swing but the fingers never move. Since the handshape carries the meaning, "
     "that is worse than nothing. Each gloss has roughly 21 candidate recordings, so candidates "
     "are scored and the best kept:")
code("""
60%   fraction of hand points detected        (handshape carries the meaning)
25%   fraction of frames with both wrists     (arm trajectory)
15%   fraction of frames with neck+shoulders  (required to normalise at all)
""")

# ── 8. State ─────────────────────────────────────────────────────────────────
doc.add_heading("8.  Current state", 1)
table(["Area", "Status"],
      [["GhSL vocabulary", "1,198 signs + 20 letters"],
       ["ASL vocabulary", "1,981 signs + 20 letters"],
       ["Avatars", "12, all VRM, 40 driven bones each, 0 unmapped"],
       ["Backend tests", "7 suites passing"],
       ["Retargeting", "Smoothing removes ~90% of jerk; no numerical failures"],
       ["Local pipeline", "Runs end to end under Docker Compose"],
       ["Cloud deployment", "Terraform written, not deployed"]],
      widths=[1.7, 4.7])

# ── 9. Limitations ───────────────────────────────────────────────────────────
doc.add_heading("9.  Limitations, stated honestly", 1)
para("Raise these yourself. A limitation you name is a boundary you understood; the same one "
     "produced by an examiner is a gap you missed.")

bullet("Both languages have 20 of 26 letters (missing a, c, l, x, y, z), and a word is only "
       "spelled when every letter exists. The wiring is complete on both the development and "
       "production paths and switches itself on when the six clips land. Google's FSboard "
       "dataset (CC BY 4.0, MediaPipe landmarks, 147 Deaf signers) was identified as the "
       "source to close it.", "Fingerspelling is off. ")
bullet("The source keypoints have z = 0 everywhere, so orientation toward or away from the "
       "camera is unrecoverable. This bounds handshape fidelity, and it particularly affects "
       "the letter pairs that GhSL distinguishes by orientation.", "The clips are 2D. ")
bullet("Around 1,200 to 2,000 signs per language cannot cover open-domain speech. "
       "Fingerspelling is the designed mitigation, which is why closing the alphabet gap "
       "matters more than adding words.", "Vocabulary is bounded. ")
bullet("Signs expire after 6 seconds, so dense speech is sampled rather than fully rendered. "
       "This is a deliberate bound on lag, not a defect, but it is a real limit on fidelity.",
       "Not everything is signed. ")
bullet("Terraform exists but has not been applied; the extension points at localhost. What has "
       "been proven is the pipeline, not the deployment.", "Nothing is deployed. ")
bullet("The GhSL dictionary is one signer's production, so regional and individual variation is "
       "unrepresented. No Deaf user study has been run.", "One signer, no user study. ")
bullet("ASL non-manual markers use the eyebrows to mark questions. The avatars carry no brow "
       "blendshapes, so that grammatical layer is currently unreachable.", "No facial grammar. ")

# ── 10. Likely questions ─────────────────────────────────────────────────────
doc.add_heading("10.  Likely examiner questions", 1)

qa("Why not just use captions?",
   "Sign languages are not visual encodings of spoken language. GhSL has its own grammar and "
   "no written form, so for a signer English captions are a second language read at a "
   "disadvantage. A signed rendering is the message rather than a translation of it. The "
   "system does use captions when they exist — as an input, because publisher text is more "
   "accurate than any ASR and carries exact timing — but the output is signing.")

qa("Is this machine translation? Where is the model?",
   "The speech recognition is a neural model, Moonshine. The text-to-gloss stage is "
   "deliberately rule-based: tokenisation, then greedy longest-match against a dictionary. It "
   "is deterministic and auditable. Given the vocabulary available, a learned model would add "
   "unpredictability without adding coverage, and an incorrect sign shown confidently is worse "
   "for a Deaf viewer than a word skipped.")

qa("Why is the vocabulary so small?",
   "It is bounded by the only public word-level GhSL keypoint dataset that exists: 1,200 signs "
   "from one signer. That is a property of the field, not a shortcut taken here. Establishing "
   "that was itself part of the work, and fingerspelling is the designed route past it.")

qa("How do you know the avatar is signing correctly?",
   "Two separate claims. Mechanically: an offline harness runs the shipped retargeting code in "
   "Node against the real avatars and clips, confirming all 40 bones resolve on all 12 rigs, "
   "no numerical failures, and about 90% of jerk removed by smoothing. Linguistically: every "
   "clip is a recording of a real signer, so no sign is invented — but no Deaf user study "
   "has been run, which is the honest limit of the claim.")

qa("What happens to a word with no sign?",
   "It is fingerspelled, if every letter is available; otherwise it is skipped and the caption "
   "line carries it. Spelling is capped at six words per utterance and 8 letters per word, "
   "because an unbounded run would queue minutes of letters and starve the signs behind it.")

qa("Why does the avatar not sign everything that is said?",
   "Signing is slower than speech: about 2.4 seconds per sign against two to three words per "
   "second. Any system that queues everything drifts further behind until it is describing "
   "something minutes old, contradicting what is on screen. Signs carry the media time they "
   "belong to and are dropped past 6 seconds, which bounds lag permanently.")

qa("What was the hardest bug?",
   "The retargeting solver used identity for the topmost driven bone's parent, ignoring the "
   "spine chain and armature rotations above it. Because error compounds down a chain, the "
   "wrist and fingers inherited all of it, and the hands looked wrung out no matter how the "
   "rest pose was written. It was found by measuring bone rotations offline rather than by "
   "looking at the render, which is the general lesson: a pose 38 degrees wrong looks "
   "plausible on screen and is obvious in a number.")

qa("Why an extension rather than an app or a website?",
   "The video is already in the browser. An extension can capture that tab's audio and draw "
   "over the page without asking the user to move content anywhere, and without any "
   "cooperation from the video platform. A website would require the publisher to integrate it.")

qa("What about privacy? You are capturing audio.",
   "Only the active tab's audio, only while capture is on, and only after an explicit user "
   "gesture that Chrome itself requires. Audio is streamed for transcription and not stored. "
   "The extension deliberately does not take the broad tabs permission, so it cannot read the "
   "addresses of the pages the user visits — tab navigation is detected from load status "
   "instead of URLs specifically to preserve that.")

qa("Can this be commercialised?",
   "The GhSL side, yes: CC BY 4.0 permits commercial use with attribution. The ASL side, not "
   "as built — WLASL is released under C-UDA, which permits academic and computational use "
   "only. Commercial ASL would require replacing that dictionary.")

qa("What would you do next?",
   "Close the alphabet from FSboard, which turns on fingerspelling for both languages and is "
   "the highest-leverage remaining change. Deploy the Terraform. Then run a Deaf user study, "
   "because every quality claim in this project is currently mechanical rather than "
   "human-validated, and that is the gap that matters most.")

# ── 11. Glossary ─────────────────────────────────────────────────────────────
doc.add_heading("11.  Glossary", 1)
table(["Term", "Meaning"],
      [["Gloss", "A written label for a sign, conventionally capitalised: THANK-YOU. Not a "
                 "translation — a name for the sign."],
       ["Keypoint", "A tracked joint position. Clips store 67: 25 body, 21 per hand."],
       ["Retargeting", "Converting recorded joint positions into bone rotations for a "
                       "particular skeleton."],
       ["Fingerspelling", "Spelling a word letter by letter with the manual alphabet, used for "
                          "names and terms with no sign."],
       ["Non-manual marker", "Grammatical information carried by the face or body rather than "
                             "the hands, such as eyebrows marking a question."],
       ["VRM", "A 3D avatar format that declares its own humanoid skeleton, so bone identity "
               "comes from the file."],
       ["MV3", "Manifest V3, Chrome's extension platform. Its service workers are ephemeral."],
       ["Offscreen document", "A hidden page an MV3 extension uses for work needing a DOM, such "
                              "as audio capture."],
       ["ONNX", "A portable model format. Moonshine ships as ONNX so it runs on CPU without a "
                "deep-learning framework."],
       ["C-UDA", "Computational Use of Data Agreement. Permits academic and computational use; "
                 "forbids commercial use."]],
      widths=[1.4, 5.0])

doc.add_paragraph()
closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = closing.add_run("Full third-party attribution is in ACKNOWLEDGEMENTS.md.\n"
                    "Architecture and repository layout are in README.md.")
r.font.size = Pt(9)
r.font.color.rgb = MUTED

import sys
out = sys.argv[1]
doc.save(out)
print("wrote", out)
